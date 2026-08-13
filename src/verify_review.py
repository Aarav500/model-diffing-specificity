"""Verify the adversarial review's factual claims against the repo's own data."""

import json
import re
from pathlib import Path

REPO = Path(__file__).resolve().parent.parent

print("=== F. diff norm claim ===")
norms = {}
for arm in ["P", "N0", "N1", "N2", "L00", "L01", "L03", "L05", "L10", "L20"]:
    p = REPO / "results" / "artifacts" / arm / "artifacts.json"
    if not p.exists():
        continue
    a = json.loads(p.read_text(encoding="utf-8"))
    vals = [v for vs in a["diff_norms_by_layer_position"].values() for v in vs]
    norms[arm] = sum(vals) / len(vals)
for k, v in norms.items():
    print(f"   {k:5s} {v:10.1f}")

lad = [("L00", 0.0), ("L01", 0.1), ("L03", 0.3), ("L05", 0.5), ("L10", 1.0), ("L20", 2.0)]
seq = [norms[a] for a, _ in lad if a in norms]
print(f"   ladder monotone rising? {all(b > a for a, b in zip(seq, seq[1:]))}  seq={[round(x) for x in seq]}")

# Does 'bigger norm = harder to find' hold ACROSS arms, as the doc implies?
print(f"   P norm {norms.get('P',0):.0f} (best accuracy) vs N0 norm {norms.get('N0',0):.0f} (no narrow objective)")
print("   -> cross-arm triage claim REFUTED by own data" if norms.get("P", 0) < norms.get("N0", 0)
      else "   -> cross-arm claim holds")

try:
    from scipy.stats import spearmanr
    acc = {"L00": .55, "L01": .35, "L03": .50, "L05": .30, "L10": .10, "L20": .05}
    xs = [norms[a] for a, _ in lad]
    ys = [acc[a] for a, _ in lad]
    r, p = spearmanr(xs, ys)
    print(f"   within mix1 family: Spearman rho={r:.2f}, p={p:.3f}")
except Exception as e:
    print("   spearman failed:", e)

print()
print("=== G. 'at most one can be right' ===")
c = json.loads((REPO / "results" / "consistency.json").read_text(encoding="utf-8"))
n0 = c.get("N0/presup", {})
print(f"   n_groups={n0.get('n_groups')} largest={n0.get('largest_group_size')} "
      f"of {n0.get('n_assertions')}  contradictory={n0.get('contradictory_pairs_exist')}")
print(f"   groups: {n0.get('groups')}")
print("   -> a 2-group split means the 4 quoted lines are NOT all mutually exclusive")

print()
print("=== C. citations missing from the doc ===")
import docx
d = docx.Document(str(REPO / "writeup" / "MATS_Nanda_Application_ModelDiffingFPR.docx"))
txt = "\n".join(p.text for p in d.paragraphs)
for t in d.tables:
    for r in t.rows:
        txt += " " + " ".join(c.text for c in r.cells)
for term in ["white-box", "Chughtai", "Engels", "Egler", "zero-delta", "black-box"]:
    print(f"   {term:14s} {'present' if term in txt else 'MISSING'}")

print()
print("=== E. deviations count claimed in doc ===")
m = re.search(r"(\w+) deviations logged", txt)
print("   doc says:", m.group(0) if m else "not found")
