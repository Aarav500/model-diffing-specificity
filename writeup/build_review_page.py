"""Assemble the review page from the converted DOCX body.

The page's job is review, not presentation: render the document faithfully and
put the mechanical facts a reviewer needs to check (word budget, figures,
placeholders) where they can be seen without scrolling.
"""

from __future__ import annotations

import re
from pathlib import Path

import docx

REPO = Path(__file__).resolve().parent.parent
DOCX = REPO / "writeup" / "MATS_Nanda_Application_ModelDiffingFPR.docx"


def facts():
    d = docx.Document(str(DOCX))
    on, n = False, 0
    for p in d.paragraphs:
        tx = p.text.strip()
        if tx == "Executive summary":
            on = True
            continue
        if tx.startswith("— end of executive summary"):
            break
        if on:
            n += len(tx.split())
    fills = sum(1 for p in d.paragraphs if any(r.font.highlight_color for r in p.runs))
    return {"words": n, "figs": len(d.inline_shapes), "fills": fills,
            "paras": len(d.paragraphs)}


CSS = """
:root{
  --paper:#FBFBF9; --ink:#16191D; --muted:#5A6169; --rule:#E2E1DC;
  --accent:#1D5C86; --critical:#A83226; --panel:#F3F3EF; --quote:#EFEFEA;
  --ok:#1E6F4F;
}
@media (prefers-color-scheme: dark){
  :root:not([data-theme="light"]){
    --paper:#14161A; --ink:#E7E5E0; --muted:#98A0A8; --rule:#2B2F35;
    --accent:#74B4DC; --critical:#E58A7B; --panel:#1B1E23; --quote:#1A1D22;
    --ok:#5FBF93;
  }
}
:root[data-theme="dark"]{
  --paper:#14161A; --ink:#E7E5E0; --muted:#98A0A8; --rule:#2B2F35;
  --accent:#74B4DC; --critical:#E58A7B; --panel:#1B1E23; --quote:#1A1D22;
  --ok:#5FBF93;
}

*{box-sizing:border-box}
body{
  margin:0; background:var(--paper); color:var(--ink);
  font-family: ui-serif, Charter, "Bitstream Charter", "Iowan Old Style", Georgia, serif;
  font-size:17px; line-height:1.62;
  -webkit-font-smoothing:antialiased;
}
.wrap{max-width:44rem; margin:0 auto; padding:0 1.5rem 6rem}

/* Review strip: the mechanical checks, visible without scrolling. */
.strip{
  position:sticky; top:0; z-index:10; background:var(--panel);
  border-bottom:1px solid var(--rule); margin-bottom:2.5rem;
}
.strip-in{
  max-width:44rem; margin:0 auto; padding:.6rem 1.5rem;
  display:flex; flex-wrap:wrap; gap:.4rem 1.5rem; align-items:baseline;
  font-family: ui-sans-serif, -apple-system, "Segoe UI", system-ui, sans-serif;
  font-size:12.5px; letter-spacing:.01em;
}
.strip b{font-variant-numeric:tabular-nums; font-weight:650}
.strip .lbl{color:var(--muted); text-transform:uppercase; letter-spacing:.07em; font-size:10.5px}
.pass{color:var(--ok)}
.critical{color:var(--critical)}

.title{
  font-family: ui-sans-serif, -apple-system, "Segoe UI", system-ui, sans-serif;
  font-size:2.05rem; line-height:1.16; font-weight:680; letter-spacing:-.022em;
  text-wrap:balance; margin:.4rem 0 .45rem;
}
.sub{
  font-family: ui-sans-serif, -apple-system, "Segoe UI", system-ui, sans-serif;
  color:var(--muted); font-size:.95rem; margin:0 0 2rem;
}

h2{
  font-family: ui-sans-serif, -apple-system, "Segoe UI", system-ui, sans-serif;
  font-size:1.32rem; font-weight:660; letter-spacing:-.014em; text-wrap:balance;
  margin:2.8rem 0 .3rem; padding-top:1.1rem; border-top:1px solid var(--rule);
}
h3{
  font-family: ui-sans-serif, -apple-system, "Segoe UI", system-ui, sans-serif;
  font-size:1.02rem; font-weight:640; letter-spacing:-.006em;
  margin:1.9rem 0 .2rem; color:var(--accent);
}
p{margin:.75rem 0}
ul{margin:.7rem 0; padding-left:1.15rem}
li{margin:.5rem 0}
strong{font-weight:660}
code{
  font-family: ui-monospace, "Cascadia Code", Consolas, monospace;
  font-size:.855em; background:var(--panel); padding:.08em .3em;
  border-radius:3px; border:1px solid var(--rule);
}
a{color:var(--accent); text-decoration-thickness:1px; text-underline-offset:2px}
a:focus-visible{outline:2px solid var(--accent); outline-offset:2px}
/* Fixed pair, deliberately not tokenised: a highlight has to read as a
   highlight on either ground, and yellow only works with dark ink on it. */
mark{background:#FDE68A; color:#16191D; padding:0 .15em}

.tw{overflow-x:auto; margin:1.1rem 0}
table{border-collapse:collapse; width:100%; font-size:.92rem;
  font-variant-numeric:tabular-nums}
th,td{border:1px solid var(--rule); padding:.45rem .6rem; text-align:left;
  vertical-align:top}
th{background:var(--panel); font-weight:640}
td:first-child{white-space:nowrap}

figure{margin:1.6rem 0 .5rem}
/* White backing is deliberate: the plots are matplotlib PNGs with dark ink on
   white, so they stay white in both themes and read as pasted-in figures --
   which is exactly what they are in the .docx. */
figure img{width:100%; height:auto; display:block;
  border:1px solid var(--rule); border-radius:2px; background:#fff}

.endmark{text-align:center; color:var(--muted); font-size:.9rem;
  letter-spacing:.03em; margin:2rem 0}

.armhead{
  font-family: ui-sans-serif, -apple-system, "Segoe UI", system-ui, sans-serif;
  font-size:.82rem; margin:1.4rem 0 .2rem; color:var(--muted);
}
.armhead strong{color:var(--ink)}
.quoted{
  font-family: ui-monospace, "Cascadia Code", Consolas, monospace;
  font-size:.78rem; line-height:1.5; margin:0; padding:0 .8rem;
  border-left:2px solid var(--rule); background:var(--quote);
  white-space:pre-wrap; color:var(--muted);
}

@media (max-width:640px){
  body{font-size:16px}
  .title{font-size:1.62rem}
}
@media (prefers-reduced-motion:reduce){*{animation:none!important;transition:none!important}}
"""


def main():
    body = (REPO / "writeup" / "_body.html").read_text(encoding="utf-8")
    f = facts()

    # First two paragraphs are the title block; promote them.
    lines = body.split("\n")
    title = re.sub(r"</?[^>]+>", "", lines[0])
    subtitle = re.sub(r"</?[^>]+>", "", lines[1])
    rest = "\n".join(lines[2:])

    words_cls = "pass" if f["words"] <= 600 else "critical"
    fills_cls = "pass" if f["fills"] == 0 else "critical"

    page = f"""<title>MATS write-up — review copy</title>
<style>{CSS}</style>
<div class="strip"><div class="strip-in">
  <span><span class="lbl">exec summary</span> <b class="{words_cls}">{f['words']}/600</b></span>
  <span><span class="lbl">figures</span> <b>{f['figs']}</b></span>
  <span><span class="lbl">placeholders</span> <b class="{fills_cls}">{f['fills']}</b></span>
  <span><span class="lbl">source</span> <b>MATS_Nanda_Application_ModelDiffingFPR.docx</b></span>
</div></div>
<div class="wrap">
  <h1 class="title">{title}</h1>
  <p class="sub">{subtitle}</p>
{rest}
</div>
"""
    out = REPO / "writeup" / "review.html"
    out.write_text(page, encoding="utf-8")
    print(f"wrote {out}  ({len(page)/1024:.0f} KB)  words={f['words']} figs={f['figs']} fills={f['fills']}")


if __name__ == "__main__":
    main()
