"""Check the document against its own pre-send checklist, before that checklist
is removed. A checklist deleted without being run is worse than no checklist.
"""

import re
from pathlib import Path

import docx

REPO = Path(__file__).resolve().parent.parent
d = docx.Document(str(REPO / "writeup" / "MATS_Nanda_Application_ModelDiffingFPR.docx"))

text = "\n".join(p.text for p in d.paragraphs)
for t in d.tables:
    for r in t.rows:
        text += "\n" + " | ".join(c.text for c in r.cells)

def check(label, ok, detail=""):
    print(f"[{'PASS' if ok else 'FAIL'}] {label}")
    if detail:
        print(f"         {detail}")

print("=== numbers discipline ===")

# The rule is "never ASSERT 97%/12% as ADL's headline". Citing it as the error
# that was caught is the opposite of a violation, so exempt sentences that
# correct it -- a blanket string ban would flag the self-correction itself.
sents_all = re.split(r"(?<=[.!?])\s+", text.replace("\n", " "))
bad = [s.strip()[:120] for s in sents_all
       if "97%" in s and not re.search(r"real:|had as|error|instead of|never", s, re.I)]
check("97% never asserted as their headline", not bad,
      f"found: {bad}" if bad else "only cited as the corrected error")

has_91 = "91%" in text
check("cites 91% (the real headline)", has_91,
      "" if has_91 else "the correct figure is missing entirely")

bad2 = re.findall(r"their controls are", text, re.I)
check('no "their controls are"', not bad2, f"found: {bad2}" if bad2 else "")

print()
print("=== grade >= 2 must not read as a refutation ===")
# Find every sentence mentioning the >=2 threshold and show it for judgement.
sents = re.split(r"(?<=[.!?])\s+", text.replace("\n", " "))
rel = [s.strip() for s in sents if "≥ 2" in s or ">= 2" in s]
print(f"  {len(rel)} sentences mention the ≥ 2 threshold:")
for s in rel:
    print(f"   - {s[:190]}")
disclaimed = any(("not comparable" in s or "not a refutation" in s.lower()) for s in rel) \
             or "not comparable to ADL" in text
check("explicitly disclaims comparability", disclaimed)

print()
print("=== the other checklist items ===")
arms = [p.text for p in d.paragraphs if re.match(r"^Arm [A-Z]\d?\d?\s", p.text)]
check("random examples pasted", len(arms) >= 10, f"{len(arms)} sampled reports embedded")
fills = [p.text for p in d.paragraphs if any(r.font.highlight_color for r in p.runs)]
check("no yellow fields remain", not fills, f"found: {fills}" if fills else "")
check("repo link present", "github.com/Aarav500/model-diffing-specificity" in text)
check("hours filled", "18 hours" in text)
