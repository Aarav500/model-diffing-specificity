"""Verify the reviewer's must-fix items landed, and re-count the word budget."""

from pathlib import Path
import docx

REPO = Path(__file__).resolve().parent.parent
d = docx.Document(str(REPO / "writeup" / "MATS_Nanda_Application_ModelDiffingFPR.docx"))

on, prose = False, 0
for p in d.paragraphs:
    t = p.text.strip()
    if t == "Executive summary":
        on = True
        continue
    if t.startswith("— end of executive summary"):
        break
    if on:
        prose += len(t.split())

# The reviewer counts the ladder table separately; report both.
tbl_words = 0
for tb in d.tables[1:]:
    for r in tb.rows:
        for c in r.cells:
            tbl_words += len(c.text.split())

txt = "\n".join(p.text for p in d.paragraphs)
for tb in d.tables:
    for r in tb.rows:
        txt += "\n" + " | ".join(c.text for c in r.cells)

print(f"exec summary prose : {prose}/600  {'OK' if prose <= 600 else 'OVER by %d' % (prose-600)}")
print(f"  + ladder table   : {prose + tbl_words} inclusive")
print()
print("must-fix 1 — the four papers")
for paper in ["Delta-Crosscoder", "AuditBench", "Model Organisms Are Leaky",
              "Cross-Architecture Diffing"]:
    print(f"   {paper:30s} {'present' if paper in txt else 'MISSING'}")
print(f"   {'field-wide framing':30s} "
      f"{'present' if 'field-wide' in txt else 'MISSING'}")
print()

# POSITIONING.md:83 BANS "cannot fail" -- Delta-Crosscoder's null is a real
# control that a broken method could fail. An earlier revision of this audit
# asserted the banned phrase, so a correct document reported MISSING. Both
# directions are now checked explicitly.
print("banned/required phrasing (POSITIONING.md:83)")
print(f"   {'BANNED cannot fail':30s} "
      f"{'PRESENT -- FIX' if 'cannot fail' in txt else 'absent (correct)'}")
print(f"   {'required zero-signal line':30s} "
      f"{'present' if 'exactly zero signal by construction' in txt else 'MISSING'}")
print()
print("must-fix 2 — registered vs exploratory separated")
print(f"   {'§8 cited':30s} {'present' if '§8' in txt else 'MISSING'}")
print(f"   {'registered test undefined':30s} "
      f"{'present' if 'undefined here' in txt else 'MISSING'}")
print(f"   {'ladder labelled exploratory':30s} "
      f"{'present' if 'exploratory' in txt else 'MISSING'}")
print(f"   {'no false pre-spec claim':30s} "
      f"{'FALSE CLAIM -- FIX' if 'both tests' in txt.lower() else 'ok'}")
print()
print("must-fix — ladder table shows all six rungs")
lad = [r.cells[0].text.strip() for r in d.tables[1].rows][1:]
print(f"   rungs in table: {lad}")
print(f"   {'six rungs':30s} {'present' if len(lad) == 6 else 'ONLY %d -- FIX' % len(lad)}")
print()
print("appendix is last (page budget)")
h1s = [p.text for p in d.paragraphs if p.style and p.style.name == "Heading 1"]
print(f"   {'appendix last':30s} "
      f"{'ok' if h1s and h1s[-1].startswith('Appendix') else 'NOT LAST -- FIX'}")
print()
print("judgment 4 — section order")
for h in [p.text for p in d.paragraphs if p.style and p.style.name == "Heading 1"]:
    print("  ", h)
print()
# Four double-encoded arrows shipped in two figure captions and a headline
# takeaway. Gate the BUILT document, not the source: the source can be clean
# while a build step reintroduces damage.
import re
MOJI = re.compile("[âÃÂ][^\x00-\x7f]")
hits = MOJI.findall(txt)
print("encoding")
print(f"   {'no mojibake in built docx':30s} "
      f"{'ok' if not hits else '%d FOUND -- FIX: %s' % (len(hits), sorted(set(map(ascii, hits))))}")
# Arrows are the specific casualty; assert the real ones survived.
print(f"   {'real arrows present':30s} "
      f"{'ok (%d)' % txt.count('→') if '→' in txt else 'NONE -- suspicious'}")
print()

# The document's credibility claim is the completeness of the deviations log,
# so the document must not contradict the log about its own size. build_docx.js
# derives this now; this check fails if that derivation ever breaks.
pre = (REPO / "PREREGISTRATION.md").read_text(encoding="utf-8")
sec = re.split(r"^#+ .*Deviations log.*$", pre, flags=re.M)[1]
sec = re.split(r"^#+ ", sec, flags=re.M)[0]
rows = [l.strip() for l in sec.splitlines() if l.strip().startswith("|")]
rows = [r for r in rows if not re.match(r"^\|[\s|:-]+\|$", r) and not r.startswith("| Date")]
WORDS = ["zero", "one", "two", "three", "four", "five", "six", "seven",
         "eight", "nine", "ten", "eleven", "twelve", "thirteen", "fourteen"]
word = WORDS[len(rows)] if len(rows) < len(WORDS) else str(len(rows))
stale = [w for w in WORDS[:15]
         if w != word and re.search(rf"\b{w} deviations\b|deviations log with {w}\b", txt, re.I)]
print("deviations count matches the log")
print(f"   {'actual rows':30s} {len(rows)} ({word})")
print(f"   {'document agrees':30s} "
      f"{'ok' if not stale else 'CONTRADICTS: says %s' % stale}")
print(f"   {'stated in document':30s} "
      f"{'yes' if re.search(rf'\b{word}\b', txt, re.I) else 'NOT STATED'}")
print()
print("unchanged framing (must NOT be smoothed)")
for phrase in ["failed to measure a false-positive rate",
               "every null I built contained a real signal",
               "What replaced it is sharper"]:
    print(f"   {'present' if phrase in txt else 'LOST'}  {phrase}")
