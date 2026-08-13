"""Faithfully convert the built .docx to HTML for review.

Reads the DOCX, not build_docx.js, so what is reviewed is what is in the file
that gets submitted. Walks body children in document order and preserves run
formatting (bold, italic, monospace, highlight), heading levels, list items,
tables, hyperlinks, and inline images (embedded as data URIs).
"""

from __future__ import annotations

import base64
import html
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn

REPO = Path(__file__).resolve().parent.parent
DOCX = REPO / "writeup" / "MATS_Nanda_Application_ModelDiffingFPR.docx"


def run_html(r, doc) -> str:
    """One <w:r>. May be an image, may be formatted text."""
    # Image?
    for blip in r.findall(".//" + qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if rid:
            part = doc.part.related_parts[rid]
            b64 = base64.b64encode(part.blob).decode()
            ctype = part.content_type or "image/png"
            return f'<img src="data:{ctype};base64,{b64}" alt="figure">'

    text = "".join(t.text or "" for t in r.findall(qn("w:t")))
    if not text:
        return ""
    out = html.escape(text)

    rPr = r.find(qn("w:rPr"))
    if rPr is not None:
        fonts = rPr.find(qn("w:rFonts"))
        if fonts is not None and (fonts.get(qn("w:ascii")) or "").startswith("Consolas"):
            out = f"<code>{out}</code>"
        if rPr.find(qn("w:b")) is not None:
            out = f"<strong>{out}</strong>"
        if rPr.find(qn("w:i")) is not None:
            out = f"<em>{out}</em>"
        hl = rPr.find(qn("w:highlight"))
        if hl is not None:
            out = f'<mark>{out}</mark>'
    return out


def para_inner(p, doc) -> str:
    """Children in document order: runs and hyperlinks."""
    parts = []
    for child in p:
        if child.tag == qn("w:r"):
            parts.append(run_html(child, doc))
        elif child.tag == qn("w:hyperlink"):
            rid = child.get(qn("r:id"))
            url = ""
            if rid and rid in doc.part.rels:
                url = doc.part.rels[rid].target_ref
            inner = "".join(run_html(r, doc) for r in child.findall(qn("w:r")))
            parts.append(f'<a href="{html.escape(url)}">{inner}</a>' if url else inner)
    return "".join(parts)


def style_of(p) -> str:
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        return ""
    st = pPr.find(qn("w:pStyle"))
    return (st.get(qn("w:val")) or "") if st is not None else ""


def convert() -> str:
    doc = docx.Document(str(DOCX))
    body = doc.element.body
    tables = iter(doc.tables)

    out: list[str] = []
    in_list = False

    def close_list():
        nonlocal in_list
        if in_list:
            out.append("</ul>")
            in_list = False

    for child in body:
        if child.tag == qn("w:tbl"):
            close_list()
            tbl = next(tables)
            out.append('<div class="tw"><table>')
            for ri, row in enumerate(tbl.rows):
                cells = []
                for c in row.cells:
                    inner = "".join(para_inner(p._p, doc) for p in c.paragraphs) or "&nbsp;"
                    cells.append(inner)
                # A row whose cells are all bold is a header row.
                hdr = ri == 0 and all("<strong>" in c for c in cells if c != "&nbsp;")
                tag = "th" if hdr else "td"
                out.append("<tr>" + "".join(f"<{tag}>{c}</{tag}>" for c in cells) + "</tr>")
            out.append("</table></div>")
            continue

        if child.tag != qn("w:p"):
            continue

        st = style_of(child)
        inner = para_inner(child, doc)
        plain = "".join(child.itertext()).strip()

        if st.startswith("Heading1"):
            close_list()
            out.append(f"<h2>{inner}</h2>")
        elif st.startswith("Heading2"):
            close_list()
            out.append(f"<h3>{inner}</h3>")
        elif st.startswith("ListParagraph"):
            if not in_list:
                out.append("<ul>")
                in_list = True
            out.append(f"<li>{inner}</li>")
        else:
            close_list()
            if not plain and "<img" not in inner:
                continue
            # Indented paragraphs are the sampled agent reports -- quoted
            # material, not body prose, and should not read as the author's voice.
            pPr = child.find(qn("w:pPr"))
            indented = pPr is not None and pPr.find(qn("w:ind")) is not None
            if "<img" in inner:
                out.append(f'<figure>{inner}</figure>')
            elif plain.startswith("— end of executive summary"):
                out.append(f'<p class="endmark">{inner}</p>')
            elif indented:
                out.append(f'<p class="quoted">{inner}</p>')
            elif plain.startswith("Arm "):
                out.append(f'<p class="armhead">{inner}</p>')
            else:
                out.append(f"<p>{inner}</p>")

    close_list()
    return "\n".join(out)


if __name__ == "__main__":
    Path(sys.argv[1] if len(sys.argv) > 1 else "writeup/_body.html").write_text(
        convert(), encoding="utf-8")
    print("ok")
