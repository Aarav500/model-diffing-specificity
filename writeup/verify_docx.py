"""Structural verification of the built .docx.

No LibreOffice on this machine, so the document cannot be rendered to PDF and
eyeballed. This checks what can be checked programmatically: heading hierarchy,
bullet lists, table shape, highlighted [FILL IN] markers, and the executive
summary word budget.
"""
import sys
import docx

path = sys.argv[1] if len(sys.argv) > 1 else "writeup/MATS_Nanda_Application_ModelDiffingFPR.docx"
d = docx.Document(path)

print("=== STRUCTURE ===")
fills = []
for p in d.paragraphs:
    tx = p.text.strip()
    if not tx:
        continue
    st = p.style.name if p.style else "?"
    hl = any(r.font.highlight_color for r in p.runs)
    if hl:
        fills.append(tx)
    marker = "[FILL]" if hl else "      "
    print(f"[{st:16s}]{marker} {tx[:94]}")

print("\n=== TABLES ===")
for n, tb in enumerate(d.tables):
    print(f"table {n}: {len(tb.rows)} x {len(tb.columns)}")
    for r in tb.rows:
        print("    " + " | ".join(c.text.strip()[:55] for c in r.cells))

print("\n=== FILL-IN MARKERS ({}) ===".format(len(fills)))
for f in fills:
    print("  -", f[:110])

on, n = False, 0
for p in d.paragraphs:
    tx = p.text.strip()
    if tx == "Executive summary":
        on = True
        continue
    if tx.startswith("— end of executive summary"):
        break
    if on:
        n += len(tx.split())
print(f"\nexec summary: {n}/600 {'OK' if n <= 600 else 'OVER'}")
print(f"total words:  {sum(len(p.text.split()) for p in d.paragraphs)}")
